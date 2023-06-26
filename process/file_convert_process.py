if 1 == 1:
    import sys
    import os

    sys.path.append(os.path.dirname(os.path.abspath(os.path.dirname(__file__))))

import time
from dtos.gui_dto import GUIDto
from common.utils import global_log_append
from timeit import default_timer as timer
from datetime import timedelta, datetime
from dtos.top_blog_detail_dto import *
from config import *
import pandas as pd
import re
from docx import Document
from common.utils import get_word_count_without_empty2

# from docx.shared import Pt, RGBColor


class FileConvert:
    def __init__(self):
        self.run_time = str(datetime.now())[0:-10].replace(":", "")

    def setGuiDto(self, guiDto: GUIDto):
        self.guiDto = guiDto

    def setLogger(self, log_msg):
        self.log_msg = log_msg

    # 워드 저장
    def sentence_to_docx(self, file_name: str, sentence: str):
        save_path = os.path.join(self.guiDto.convert_path, f"파일 변환 {self.run_time}")

        if os.path.isdir(save_path) == False:
            os.mkdir(save_path)
        else:
            pass

        sentence_count = get_word_count_without_empty2(sentence)
        sentence_docx = os.path.join(save_path, f"{file_name}_{sentence_count}.docx")

        doc = Document()
        doc.add_paragraph(sentence)
        doc.save(sentence_docx)
        self.log_msg.emit(f"{file_name}.docx 저장 완료")

    # 메모장 저장
    def sentence_to_txt(self, file_name: str, sentence: str):
        save_path = os.path.join(self.guiDto.convert_path, f"파일 변환 {self.run_time}")

        if os.path.isdir(save_path) == False:
            os.mkdir(save_path)
        else:
            pass

        sentence_txt = os.path.join(save_path, f"{file_name}.txt")

        with open(sentence_txt, "w", encoding="utf-8") as f:
            f.write(sentence)

        self.log_msg.emit(f"{file_name}.txt 저장 완료")

    # 포스터 워드 저장
    def poster_sentence_to_docx(self, file_name: str, sentence: str):
        save_path = os.path.join(self.guiDto.convert_path, f"포스터용 변환 {self.run_time}")

        if os.path.isdir(save_path) == False:
            os.mkdir(save_path)
        else:
            pass

        file_path = os.path.join(save_path, file_name)

        if os.path.isdir(file_path) == False:
            os.mkdir(file_path)
        else:
            pass

        sentence_docx = os.path.join(file_path, f"{file_name}.docx")
        doc = Document()
        doc.add_paragraph(sentence)
        doc.save(sentence_docx)

        self.log_msg.emit(f"{file_name}.docx 저장 완료")

    # 포스터 메모장 저장
    def poster_sentence_to_txt(self, file_name: str, sentence: str):
        save_path = os.path.join(self.guiDto.convert_path, f"포스터용 변환 {self.run_time}")

        if os.path.isdir(save_path) == False:
            os.mkdir(save_path)
        else:
            pass

        file_path = os.path.join(save_path, file_name)

        if os.path.isdir(file_path) == False:
            os.mkdir(file_path)
        else:
            pass

        sentence_txt = os.path.join(file_path, f"{file_name}.txt")

        with open(sentence_txt, "w", encoding="utf-8") as f:
            f.write(sentence)

        self.log_msg.emit(f"{file_name}.txt 저장 완료")

    def get_sentence_from_file(self, file_path: str):
        sentence = ""

        # docx 파일인 경우
        if file_path.rfind(".docx") > -1:
            doc = Document(file_path)
            all_text = []
            for para in doc.paragraphs:
                all_text.append(para.text + "\n")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text.append(cell.text)
            all_text = " ".join(all_text)
            sentence = all_text

        # txt 파일인 경우
        elif file_path.rfind(".txt") > -1:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
                sentence = text
        else:
            pass

        sentence = sentence.replace("\n ", "\n")

        return sentence

    # 전체작업 시작
    def work_start(self):
        file: str
        for i, file in enumerate(self.guiDto.convert_list):
            file_path = os.path.join(self.guiDto.convert_path, file)
            file_format = file[file.rfind(".") :]

            # 파일 유효성 검사
            if os.path.isfile(file_path):
                print(file_path)
            else:
                continue

            # 파일에서 문자열 획득
            original_sentence = self.get_sentence_from_file(file_path)

            if not self.guiDto.poster_option:
                # 파일 저장
                if self.guiDto.convert_format == "txt":
                    self.sentence_to_txt(file.rstrip(file_format), original_sentence)
                elif self.guiDto.convert_format == "docx":
                    self.sentence_to_docx(file.rstrip(file_format), original_sentence)

            elif self.guiDto.poster_option:
                # 티스토리 포스터용으로 변환하기
                if self.guiDto.convert_format == "txt":
                    self.poster_sentence_to_txt(file.rstrip(file_format), original_sentence)
                elif self.guiDto.convert_format == "docx":
                    self.poster_sentence_to_docx(file.rstrip(file_format), original_sentence)


if __name__ == "__main__":
    converter = FileConvert()
    converter.work_start()
