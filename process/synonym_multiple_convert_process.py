if 1 == 1:
    import sys
    import os

    sys.path.append(os.path.dirname(os.path.abspath(os.path.dirname(__file__))))

import time
from dtos.gui_dto import GUIDto
from common.utils import global_log_append
from timeit import default_timer as timer
from datetime import timedelta, datetime
from dtos.top_blog_detail_dto import *
from config import *
import pandas as pd
import re
from docx import Document
from docx.shared import RGBColor
from docx.enum import text
from features.convert_sentence import (
    convert_from_db,
    shuffle_sentence,
    replace_keyword,
)
import random
from docx.text.paragraph import Paragraph
from pathlib import Path


class SynonymMultipleConvert:
    def __init__(self):
        self.run_time = str(datetime.now())[0:-10].replace(":", "")

    def setGuiDto(self, guiDto: GUIDto):
        self.guiDto = guiDto

    def setLogger(self, log_msg):
        self.log_msg = log_msg

    # 워드 저장
    def sentence_to_docx(
        self,
        file_name: str,
        header: str,
        footer: str,
        dict_sentence: dict,
        used_idx_list: list,
        limit="",
    ):
        print(dict_sentence)

        save_path = os.path.join(self.guiDto.convert_path, f"유의어 변환 {self.run_time}")

        if os.path.isdir(save_path) == False:
            os.mkdir(save_path)
        else:
            pass

        if limit == "":
            sentence_docx = os.path.join(save_path, f"{file_name}.docx")
        else:
            limit = str(limit)
            sentence_docx = os.path.join(
                save_path, f"{file_name}_{limit.zfill(2)}.docx"
            )

        doc = Document()
        paragraph = doc.add_paragraph()

        # 머리말 입력
        if header:
            for header_word in header:
                run = paragraph.add_run(header_word)
                font = run.font
                font.color.rgb = RGBColor(0, 0, 255)  # 빨간색으로 설정

            paragraph = doc.add_paragraph()
            paragraph = doc.add_paragraph()

        # 본문입력 # 여기까지 넣고난 뒤에 섞어야함.
        for sentence_i in dict_sentence.keys():
            word = dict_sentence[sentence_i]
            if word == "\n":
                paragraph = doc.add_paragraph()
            else:
                run = paragraph.add_run(word)

            if sentence_i in used_idx_list:
                font = run.font
                font.color.rgb = RGBColor(255, 0, 0)  # 빨간색으로 설정

        # 맺음말 입력
        if footer:
            paragraph = doc.add_paragraph()
            for footer_word in footer:
                run = paragraph.add_run(footer_word)
                font = run.font
                font.color.rgb = RGBColor(0, 0, 255)  # 빨간색으로 설정

        doc.save(sentence_docx)

        if limit == "":
            self.log_msg.emit(f"{file_name}.docx 저장 완료")
        else:
            limit = str(limit)
            self.log_msg.emit(f"{file_name}_{limit.zfill(2)}.docx 저장 완료")

    def get_sentence_from_file(self, file_path: str):
        sentence = ""

        # docx 파일인 경우
        if file_path.rfind(".docx") > -1:
            doc = Document(file_path)
            all_text = []
            for para in doc.paragraphs:
                all_text.append(para.text + "\n")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text.append(cell.text)
            all_text = " ".join(all_text)
            sentence = all_text

        # txt 파일인 경우
        elif file_path.rfind(".txt") > -1:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
                sentence = text
        else:
            pass

        sentence = sentence.replace("\n ", "\n")

        return sentence

    # 전체작업 시작
    def work_start(self):
        print(f"converter: work_start {self.run_time}")

        file: str
        for i, file in enumerate(self.guiDto.convert_list):
            file_path = os.path.join(self.guiDto.convert_path, file)
            file_format = file[file.rfind(".") :]

            # 파일 유효성 검사
            if os.path.isfile(file_path):
                print(file_path)
            else:
                continue

            # 파일에서 문자열 획득
            original_sentence = self.get_sentence_from_file(file_path)
            file_name = Path(file_path).stem

            # 횟수 제한 기능
            for limit in range(1, self.guiDto.synonym_convert_limit + 1):
                # 문단 랜덤 섞기 체크 시

                if self.guiDto.shuffle_paragraphs_check:
                    original_sentence = shuffle_sentence(original_sentence)

                dict_sentence = {}
                used_idx_list = []

                # 문자열 변환
                dict_sentence, used_idx_list = convert_from_db(
                    original_sentence,
                    file_name,
                    self.guiDto.df_two_way,
                    self.guiDto.df_one_way,
                )

                print(dict_sentence)

                # 머리글 삽입
                header = ""
                if self.guiDto.header_check:
                    header_topic = self.guiDto.header_topic
                    saved_data_header = get_save_data_HEADER()
                    header: str = random.choice(saved_data_header[header_topic])
                    header = replace_keyword(header, file.rstrip(file_format))

                # 맺음말 삽입
                footer = ""
                if self.guiDto.footer_check:
                    footer_topic = self.guiDto.footer_topic
                    saved_data_footer = get_save_data_FOOTER()
                    footer: str = random.choice(saved_data_footer[footer_topic])
                    footer = replace_keyword(footer, file.rstrip(file_format))

                # 문자열 파일 저장
                file_name = file.rstrip(file_format)

                self.sentence_to_docx(
                    file_name,
                    header,
                    footer,
                    dict_sentence,
                    used_idx_list,
                    limit,
                )


if __name__ == "__main__":
    guiDto = GUIDto()
    guiDto.synonym_convert_limit = 1
    guiDto.convert_list = [
        r"C:\consolework\tistory-convert\__test__\테스트_230428\감 효능.docx"
    ]

    file_path = r"C:\consolework\tistory-convert\excel\유의어db_정렬변경0416.xlsx"

    # two_way_columns = self.two_way_data_type()
    try:
        df_two_way: pd.DataFrame = pd.read_excel(
            file_path, sheet_name="양방향", keep_default_na=""
        )
        df_two_way = df_two_way.astype(str)
        # self.df_two_way = self.df_two_way.loc[:, list(two_way_columns.keys())]
    except Exception as e:
        print(e)

    try:
        df_one_way = pd.read_excel(
            file_path,
            converters={"before": str, "after": str},
            sheet_name="일방향",
            keep_default_na="",
        )
        df_one_way = df_one_way.loc[:, ["before", "after"]]
    except Exception as e:
        print(e)

    guiDto.df_two_way = df_two_way
    guiDto.df_one_way = df_one_way

    converter = SynonymMultipleConvert()
    converter.setGuiDto(guiDto=guiDto)
    converter.work_start()
